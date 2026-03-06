"""
Script pour générer le rapport de projet MauriBus au format Word (.docx)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '059669')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_placeholder_box(doc, label="[ Insérer capture d'écran ici ]"):
    """Adds a gray placeholder box for images/diagrams."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run.font.italic = True
    # Add a visible border via table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    # Set cell height (approximate via paragraph)
    for _ in range(8):
        cell.add_paragraph()
    set_cell_bg(cell, 'F3F4F6')
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run2.font.italic = True
    doc.add_paragraph()

# ─── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ─── Page de garde ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSITÉ DE NOUAKCHOTT")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Faculté des Sciences et Techniques")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Département Informatique")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

add_horizontal_rule(doc)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE PROJET DE FIN D'ÉTUDES")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MauriBus")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Système de Gestion Intelligente du Transport Urbain en Mauritanie")
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# Placeholder logo
add_placeholder_box(doc, "[ Logo MauriBus ]")

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Réalisé par :")
run.bold = True
run.font.size = Pt(12)

students = [
    "Lalle Vatma Mamine",
    "Mariem Didi",
    "Fatimetou Cheik ould meny",
    "Oumkelthoum Mouhamedou",
]
for s in students:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(s)
    run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Encadrant :")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prof. Aiche")
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Année Universitaire 2024 – 2025")
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ─── Remerciements ─────────────────────────────────────────────────────────────
h = doc.add_heading('Remerciements', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph(
    "Avant toute chose, nous tenons à exprimer notre profonde gratitude envers toutes les "
    "personnes qui ont contribué, de près ou de loin, à la réalisation de ce projet."
)
doc.add_paragraph(
    "Nous adressons nos sincères remerciements à notre encadrant, le Professeur Aiche, "
    "pour son soutien constant, ses précieux conseils et sa disponibilité tout au long de "
    "ce travail. Sa rigueur scientifique et son expertise ont été une source d'inspiration "
    "et un guide indispensable pour la réussite de ce projet."
)
doc.add_paragraph(
    "Nous remercions également l'ensemble du corps enseignant du Département Informatique "
    "de la Faculté des Sciences et Techniques de l'Université de Nouakchott pour la qualité "
    "de la formation dispensée et les connaissances transmises durant notre cursus."
)
doc.add_paragraph(
    "Nos remerciements vont aussi à nos familles et amis pour leur soutien moral et leurs "
    "encouragements constants qui nous ont permis de persévérer face aux difficultés "
    "rencontrées au cours de ce projet."
)
doc.add_paragraph(
    "Enfin, nous remercions toutes les personnes qui ont accepté de tester l'application "
    "et de nous faire part de leurs retours, contribuant ainsi à l'amélioration de MauriBus."
)

doc.add_page_break()

# ─── Résumé ────────────────────────────────────────────────────────────────────
h = doc.add_heading('Résumé', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph(
    "MauriBus est une solution numérique complète dédiée à la gestion intelligente du "
    "transport urbain en Mauritanie. Le système se compose de trois composants principaux : "
    "une interface d'administration web, une application mobile pour les chauffeurs et une "
    "application mobile pour les citoyens."
)
doc.add_paragraph(
    "L'interface d'administration, développée en React.js, permet aux administrateurs de "
    "gérer l'intégralité du réseau de transport : chauffeurs, bus, lignes, trajets, "
    "paiements et notifications. Elle intègre également une carte en temps réel affichant "
    "les positions GPS de tous les bus actifs."
)
doc.add_paragraph(
    "L'application mobile chauffeur, développée avec React Native (Expo), permet aux "
    "conducteurs de gérer leurs trajets, d'envoyer leur position GPS en continu, de "
    "valider les tickets des passagers via QR code et de soumettre des signalements "
    "d'incidents."
)
doc.add_paragraph(
    "L'application mobile citoyenne offre aux usagers la possibilité de consulter les "
    "lignes disponibles, de visualiser les bus en temps réel, de payer leurs trajets "
    "via un wallet électronique intégré, d'obtenir un ticket QR et de noter les chauffeurs."
)
doc.add_paragraph(
    "Le backend est développé en Django (Python) avec Django REST Framework, "
    "utilisant une authentification JWT personnalisée pour les trois types d'utilisateurs "
    "(administrateurs, chauffeurs, citoyens). La base de données SQLite est utilisée "
    "pour le développement, avec une compatibilité PostgreSQL pour la production."
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Mots-clés : ")
run.bold = True
p.add_run("Transport urbain, Application mobile, React Native, Django, GPS, Wallet électronique, QR Code, Mauritanie.")

doc.add_page_break()

# ─── Abstract ──────────────────────────────────────────────────────────────────
h = doc.add_heading('Abstract', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph(
    "MauriBus is a comprehensive digital solution dedicated to intelligent urban transport "
    "management in Mauritania. The system consists of three main components: a web "
    "administration interface, a mobile application for drivers, and a mobile application "
    "for citizens."
)
doc.add_paragraph(
    "The web administration interface, developed in React.js, enables administrators to "
    "manage the entire transport network: drivers, buses, routes, trips, payments, and "
    "notifications. It also features a real-time map displaying the GPS positions of all "
    "active buses."
)
doc.add_paragraph(
    "The driver mobile application, developed with React Native (Expo), allows drivers "
    "to manage their trips, continuously send GPS positions, validate passenger tickets "
    "via QR code scanning, and submit incident reports."
)
doc.add_paragraph(
    "The citizen mobile application provides users with the ability to consult available "
    "bus lines, view buses in real-time, pay for trips through an integrated electronic "
    "wallet, obtain a QR ticket, and rate drivers."
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.bold = True
p.add_run("Urban transport, Mobile application, React Native, Django, GPS, Electronic wallet, QR Code, Mauritania.")

doc.add_page_break()

# ─── Table des matières ────────────────────────────────────────────────────────
h = doc.add_heading('Table des Matières', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

toc_items = [
    ("Remerciements", ""),
    ("Résumé / Abstract", ""),
    ("Table des Matières", ""),
    ("Liste des Figures", ""),
    ("Liste des Tableaux", ""),
    ("Introduction Générale", ""),
    ("Chapitre 1 : Contexte et Problématique", ""),
    ("    1.1  Contexte du projet", ""),
    ("    1.2  Problématique", ""),
    ("    1.3  Objectifs", ""),
    ("    1.4  Cahier des charges", ""),
    ("Chapitre 2 : État de l'Art", ""),
    ("    2.1  Solutions existantes", ""),
    ("    2.2  Comparaison des solutions", ""),
    ("    2.3  Positionnement de MauriBus", ""),
    ("Chapitre 3 : Analyse et Conception", ""),
    ("    3.1  Architecture globale", ""),
    ("    3.2  Diagramme de cas d'utilisation", ""),
    ("    3.3  Diagramme de classes", ""),
    ("    3.4  Modèle de données", ""),
    ("Chapitre 4 : Implémentation", ""),
    ("    4.1  Environnement de développement", ""),
    ("    4.2  Backend Django", ""),
    ("    4.3  Interface d'administration Web", ""),
    ("    4.4  Application mobile Chauffeur", ""),
    ("    4.5  Application mobile Citoyen", ""),
    ("    4.6  Système GPS et carte en temps réel", ""),
    ("    4.7  Wallet électronique et paiements", ""),
    ("    4.8  Système de tickets QR", ""),
    ("Chapitre 5 : Tests et Validation", ""),
    ("    5.1  Tests fonctionnels", ""),
    ("    5.2  Tests des API", ""),
    ("    5.3  Tests des applications mobiles", ""),
    ("Chapitre 6 : Résultats et Discussion", ""),
    ("    6.1  Présentation de l'application", ""),
    ("    6.2  Difficultés rencontrées", ""),
    ("    6.3  Perspectives d'amélioration", ""),
    ("Conclusion Générale", ""),
    ("Bibliographie / Webographie", ""),
    ("Annexes", ""),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("    "):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ─── Liste des figures ─────────────────────────────────────────────────────────
h = doc.add_heading('Liste des Figures', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
doc.add_paragraph("Figure 1 : Architecture globale du système MauriBus")
doc.add_paragraph("Figure 2 : Diagramme de cas d'utilisation")
doc.add_paragraph("Figure 3 : Diagramme de classes")
doc.add_paragraph("Figure 4 : Interface d'administration — Tableau de bord")
doc.add_paragraph("Figure 5 : Interface d'administration — Carte GPS en temps réel")
doc.add_paragraph("Figure 6 : Interface d'administration — Gestion des lignes")
doc.add_paragraph("Figure 7 : Application chauffeur — Tableau de bord")
doc.add_paragraph("Figure 8 : Application chauffeur — Écran de trajet actif")
doc.add_paragraph("Figure 9 : Application chauffeur — Scanner QR ticket")
doc.add_paragraph("Figure 10 : Application citoyenne — Tableau de bord avec carte")
doc.add_paragraph("Figure 11 : Application citoyenne — Écran wallet")
doc.add_paragraph("Figure 12 : Application citoyenne — Ticket QR généré")

doc.add_page_break()

# ─── Liste des tableaux ────────────────────────────────────────────────────────
h = doc.add_heading('Liste des Tableaux', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
doc.add_paragraph("Tableau 1 : Technologies utilisées")
doc.add_paragraph("Tableau 2 : Description des acteurs du système")
doc.add_paragraph("Tableau 3 : Endpoints API Backend")
doc.add_paragraph("Tableau 4 : Comparaison des solutions de transport")
doc.add_paragraph("Tableau 5 : Résultats des tests fonctionnels")

doc.add_page_break()

# ─── Introduction Générale ─────────────────────────────────────────────────────
h = doc.add_heading('Introduction Générale', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph(
    "Le transport urbain constitue un enjeu fondamental pour le développement économique "
    "et social de toute ville. En Mauritanie, et plus particulièrement dans la capitale "
    "Nouakchott, le secteur des transports en commun fait face à de nombreux défis : "
    "absence de système de suivi en temps réel, difficultés de gestion des flottes de bus, "
    "manque de transparence dans les paiements et insuffisance d'information pour les usagers."
)
doc.add_paragraph(
    "Face à ces problèmes, nous avons conçu et développé MauriBus, une solution numérique "
    "intégrée visant à moderniser et à digitaliser la gestion du transport urbain. "
    "Ce projet s'inscrit dans une démarche de transformation numérique du secteur public "
    "mauritanien, en exploitant les technologies mobiles et web les plus récentes."
)
doc.add_paragraph(
    "MauriBus propose une approche à trois niveaux : une interface d'administration "
    "centralisée pour les gestionnaires du réseau, une application mobile dédiée aux "
    "chauffeurs pour le suivi de leurs activités, et une application mobile pour les "
    "citoyens facilitant l'accès aux services de transport et le paiement électronique."
)
doc.add_paragraph(
    "Ce rapport présente l'ensemble du processus de développement de MauriBus, depuis "
    "l'analyse des besoins jusqu'à l'implémentation et les tests, en passant par la "
    "conception architecturale du système. Il est structuré en six chapitres couvrant "
    "respectivement : le contexte et la problématique, l'état de l'art, l'analyse et "
    "la conception, l'implémentation technique, les tests et validation, et enfin les "
    "résultats obtenus et les perspectives d'amélioration."
)

doc.add_page_break()

# ─── Chapitre 1 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 1 : Contexte et Problématique', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('1.1  Contexte du projet', level=2)
doc.add_paragraph(
    "La Mauritanie, pays en voie de développement situé en Afrique de l'Ouest, connaît "
    "une urbanisation croissante. Nouakchott, sa capitale, compte plus d'un million "
    "d'habitants et dispose d'un réseau de transport en commun géré de manière traditionnelle, "
    "sans outils numériques modernes. Les opérateurs de transport peinent à coordonner "
    "leurs flottes, les chauffeurs manquent d'outils de gestion et les usagers n'ont "
    "accès à aucune information fiable sur les horaires ou les itinéraires des bus."
)
doc.add_paragraph(
    "C'est dans ce contexte que s'inscrit le projet MauriBus. Initié dans le cadre de "
    "notre formation en informatique, il vise à apporter une réponse concrète et "
    "technologique aux lacunes du secteur des transports en commun mauritanien."
)

doc.add_heading('1.2  Problématique', level=2)
doc.add_paragraph(
    "La gestion du transport urbain à Nouakchott souffre de plusieurs problèmes majeurs :"
)
problems = [
    "Absence de suivi en temps réel des bus : les gestionnaires ne savent pas où se trouvent leurs véhicules.",
    "Opacité des paiements : les revenus des trajets sont difficiles à contrôler et à tracer.",
    "Manque d'information pour les usagers : pas d'application permettant de consulter les lignes, horaires ou tarifs.",
    "Gestion papier des tickets : pas de système numérique de billetterie.",
    "Difficulté de coordination : aucun canal de communication structuré entre administrateurs et chauffeurs.",
    "Absence de système de notation des chauffeurs : impossible d'évaluer la qualité du service.",
]
for pb in problems:
    p = doc.add_paragraph(pb, style='List Bullet')

doc.add_heading('1.3  Objectifs', level=2)
doc.add_paragraph("Le projet MauriBus a pour objectifs principaux :")
objectives = [
    "Centraliser la gestion du réseau de transport dans une interface d'administration unique.",
    "Permettre le suivi GPS en temps réel de tous les bus.",
    "Digitaliser les paiements via un wallet électronique intégré.",
    "Fournir aux usagers une application mobile moderne pour accéder aux services de transport.",
    "Automatiser la validation des tickets via QR code.",
    "Améliorer la communication entre administrateurs et chauffeurs.",
    "Générer des statistiques et rapports d'activité automatiques.",
]
for obj in objectives:
    p = doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.4  Cahier des charges', level=2)
doc.add_paragraph("Le système doit satisfaire aux exigences fonctionnelles et non-fonctionnelles suivantes :")

doc.add_heading('Exigences fonctionnelles', level=3)
func_req = [
    "Authentification sécurisée pour les trois types d'utilisateurs (Admin, Chauffeur, Citoyen).",
    "Gestion CRUD complète des entités (chauffeurs, bus, lignes, arrêts, trajets).",
    "Suivi GPS en temps réel avec affichage cartographique.",
    "Système de paiement électronique avec historique des transactions.",
    "Génération et validation de tickets QR.",
    "Système de notifications push.",
    "Tableau de bord avec statistiques et indicateurs clés.",
    "Gestion des signalements d'incidents.",
    "Système de notation des chauffeurs.",
]
for r in func_req:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Exigences non-fonctionnelles', level=3)
nonfunc_req = [
    "Performance : réponse des API < 500ms.",
    "Sécurité : authentification JWT, mots de passe hachés (PBKDF2).",
    "Disponibilité : compatible iOS et Android.",
    "Scalabilité : architecture permettant l'ajout de nouvelles lignes et chauffeurs.",
    "Ergonomie : interface intuitive adaptée au contexte mauritanien.",
]
for r in nonfunc_req:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ─── Chapitre 2 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 2 : État de l\'Art', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('2.1  Solutions existantes', level=2)
doc.add_paragraph(
    "Plusieurs solutions de gestion de transport urbain existent à travers le monde. "
    "Nous en présentons ici les plus pertinentes pour notre contexte."
)

doc.add_heading('Google Maps / Transit', level=3)
doc.add_paragraph(
    "Google Maps propose une fonctionnalité de navigation en transport en commun. "
    "Cependant, elle nécessite que les opérateurs publient leurs données au format GTFS "
    "(General Transit Feed Specification) — ce qui n'est pas le cas en Mauritanie. "
    "De plus, elle ne propose aucun outil de gestion interne pour les opérateurs."
)

doc.add_heading('Moovit', level=3)
doc.add_paragraph(
    "Moovit est une application de planification de trajets en transport en commun "
    "disponible dans de nombreuses villes. Comme Google Maps, elle dépend de données "
    "GTFS et ne propose pas d'outil de gestion de flotte ni de paiement intégré."
)

doc.add_heading('Padam Mobility / Karhoo', level=3)
doc.add_paragraph(
    "Ces solutions B2B proposent des outils complets de gestion de flotte et de "
    "demand-responsive transport. Elles sont cependant conçues pour des marchés "
    "développés, avec des coûts de licence élevés et peu adaptées au contexte "
    "africain ou aux petits opérateurs."
)

doc.add_heading('2.2  Comparaison des solutions', level=2)

# Comparison table
table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'
headers = ['Critère', 'Google Maps', 'Moovit', 'Padam', 'MauriBus']
for i, h_text in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h_text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, 'D1FAE5')

rows_data = [
    ['Suivi GPS temps réel', '✓', '✓', '✓', '✓'],
    ['Gestion de flotte', '✗', '✗', '✓', '✓'],
    ['Paiement intégré', '✗', '✗', '✓', '✓'],
    ['Adapté à la Mauritanie', '✗', '✗', '✗', '✓'],
]
for i, row_data in enumerate(rows_data):
    row = table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

doc.add_heading('2.3  Positionnement de MauriBus', level=2)
doc.add_paragraph(
    "MauriBus se distingue des solutions existantes par son adaptation spécifique au "
    "contexte mauritanien. Contrairement aux solutions internationales, MauriBus intègre :"
)
distinctions = [
    "Un système de paiement via Bankily, la principale solution de paiement mobile en Mauritanie.",
    "Une interface disponible en langue française et adaptée aux pratiques locales.",
    "Un système de gestion complète accessible aux opérateurs de petite et moyenne taille.",
    "Un coût de déploiement minimal, adapté aux contraintes budgétaires du secteur public mauritanien.",
    "Une architecture open source, permettant des adaptations futures.",
]
for d in distinctions:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# ─── Chapitre 3 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 3 : Analyse et Conception', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('3.1  Architecture globale', level=2)
doc.add_paragraph(
    "MauriBus repose sur une architecture trois-tiers classique, déclinée en trois "
    "composants frontaux et un serveur backend central :"
)
arch = [
    "Interface d'administration (React.js + Vite) — Port 5173 en développement",
    "Application mobile chauffeur (React Native + Expo)",
    "Application mobile citoyenne (React Native + Expo)",
    "Backend API (Django 4.2 + Django REST Framework) — Port 8000",
    "Base de données (SQLite en développement, PostgreSQL en production)",
]
for a in arch:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph(
    "Les trois composants frontaux communiquent avec le backend via des API REST "
    "sécurisées par JWT. Chaque type d'utilisateur dispose d'un token JWT distinct "
    "(admin_id, driver_id ou citizen_id dans le payload) et d'une classe "
    "d'authentification spécifique côté Django."
)

add_placeholder_box(doc, "[ Figure 1 : Schéma d'architecture globale de MauriBus ]")

doc.add_heading('3.2  Diagramme de cas d\'utilisation', level=2)
doc.add_paragraph(
    "Le diagramme de cas d'utilisation identifie cinq acteurs principaux dans le système :"
)
actors = [
    "Super Administrateur : gère les administrateurs, consulte les statistiques globales.",
    "Administrateur : gère chauffeurs, bus, lignes, trajets, notifications et paiements.",
    "Modérateur : consulte les données en lecture seule.",
    "Chauffeur : gère ses trajets, envoie sa position GPS, valide les tickets, soumet des signalements.",
    "Citoyen : consulte les lignes, visualise les bus en temps réel, paye ses trajets, gère son wallet.",
]
for a in actors:
    doc.add_paragraph(a, style='List Bullet')

add_placeholder_box(doc, "[ Figure 2 : Diagramme de cas d'utilisation ]")

doc.add_heading('3.3  Diagramme de classes', level=2)
doc.add_paragraph(
    "Le modèle de données de MauriBus comprend 16 classes principales, réparties en "
    "quatre domaines fonctionnels :"
)
domains = [
    "Domaine Administration : Admin",
    "Domaine Transport : Driver, Bus, Line, Stop, Trip, GPSPosition, Report, Notification, Payment",
    "Domaine Citoyen : CitizenUser, WalletCitizen, WalletTransaction, CitizenTrip, DriverRating",
    "Domaine Finance : WalletDriver, AppWallet, AppTransaction",
]
for d in domains:
    doc.add_paragraph(d, style='List Bullet')

add_placeholder_box(doc, "[ Figure 3 : Diagramme de classes UML ]")

doc.add_heading('3.4  Modèle de données', level=2)
doc.add_paragraph(
    "Le modèle de données utilise des UUID (Universally Unique Identifiers) comme clés "
    "primaires pour toutes les entités métier, garantissant l'unicité globale et "
    "facilitant une future migration vers une architecture distribuée. Les relations "
    "principales sont :"
)
relations = [
    "Bus ↔ Driver : relation OneToOne (un bus est assigné à un chauffeur maximum)",
    "Bus ↔ Line : relation ManyToMany (un bus peut circuler sur plusieurs lignes)",
    "Stop → Line : relation ManyToOne (une ligne a plusieurs arrêts ordonnés)",
    "Trip → Driver, Bus, Line : le trajet lie le chauffeur, le bus et la ligne",
    "GPSPosition → Bus, Trip : positions enregistrées par bus et par trajet",
    "CitizenTrip → CitizenUser, Line, Trip : historique des trajets payés",
    "WalletCitizen ↔ CitizenUser : wallet OneToOne créé automatiquement à l'inscription",
]
for r in relations:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ─── Chapitre 4 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 4 : Implémentation', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('4.1  Environnement de développement', level=2)
doc.add_paragraph(
    "Le développement de MauriBus a été réalisé dans l'environnement suivant :"
)

# Technologies table
tech_table = doc.add_table(rows=13, cols=3)
tech_table.style = 'Table Grid'
tech_headers = ['Composant', 'Technologie', 'Version']
for i, h_text in enumerate(tech_headers):
    cell = tech_table.cell(0, i)
    cell.text = h_text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, 'D1FAE5')

tech_data = [
    ['Backend', 'Python / Django', '3.11 / 4.2'],
    ['API REST', 'Django REST Framework', '3.14'],
    ['Authentification', 'JWT (PyJWT)', '2.8'],
    ['Base de données', 'SQLite / PostgreSQL', '3.x / 15'],
    ['Interface admin', 'React.js + Vite', '18 / 5'],
    ['Composants UI', 'Tailwind CSS', '3.4'],
    ['Icônes', 'React Icons', '5.x'],
    ['App mobile', 'React Native (Expo)', 'SDK 54'],
    ['Navigation', 'React Navigation', '6'],
    ['Cartes', 'React Leaflet / MapView', '—'],
    ['QR Code', 'react-native-qrcode-svg', '6.3'],
    ['Caméra', 'expo-camera', '16'],
]
for i, row_data in enumerate(tech_data):
    row = tech_table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

doc.add_heading('4.2  Backend Django', level=2)
doc.add_paragraph(
    "Le backend de MauriBus est un serveur Django exposant une API REST. Son organisation "
    "suit l'architecture suivante :"
)

doc.add_heading('Structure du backend', level=3)
backend_struct = [
    "core/models/ — Modèles Django : admin.py, driver.py, bus.py, citizen.py",
    "core/views/ — Vues API : auth_views.py, driver_views.py, citizen_views.py, admin_extra_views.py, viewsets.py",
    "core/authentication.py — Classes d'authentification JWT (Admin, Driver, Citizen)",
    "core/urls.py — Routage des endpoints",
    "mauribus_project/settings.py — Configuration Django",
]
for s in backend_struct:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Authentification JWT', level=3)
doc.add_paragraph(
    "L'authentification repose sur des tokens JWT personnalisés. Contrairement à l'approche "
    "standard Django (AbstractUser), MauriBus utilise des modèles indépendants (Admin, "
    "Driver, CitizenUser) avec des classes d'authentification DRF dédiées. Chaque token "
    "contient dans son payload un identifiant spécifique (admin_id, driver_id, citizen_id) "
    "permettant au backend d'identifier le type et l'identité de l'appelant."
)

doc.add_heading('Endpoints API', level=3)
doc.add_paragraph(
    "Le backend expose plus de 40 endpoints organisés en préfixes :"
)
endpoints = [
    "/api/v1/auth/ — Authentification administrateurs",
    "/api/v1/driver/ — Actions chauffeurs (GPS, trajets, tickets, signalements)",
    "/api/v1/citizen/ — Actions citoyens (inscription, wallet, paiements, lignes)",
    "/api/v1/admin/ — Gestion CRUD (viewsets) et fonctions avancées",
]
for ep in endpoints:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_heading('4.3  Interface d\'administration Web', level=2)
doc.add_paragraph(
    "L'interface d'administration est une Single Page Application (SPA) développée "
    "en React.js avec Vite comme bundler. Elle communique avec le backend via Axios, "
    "avec un proxy Vite rewritant /api/ vers /api/v1/ sur le serveur Django."
)
doc.add_paragraph(
    "Les principales pages de l'interface sont :"
)
admin_pages = [
    "Tableau de bord — statistiques globales, indicateurs clés",
    "Gestion des chauffeurs — CRUD, suspension, visualisation des notes",
    "Gestion des bus — CRUD, assignation des chauffeurs, statut",
    "Gestion des lignes — CRUD, gestion des arrêts, tarifs",
    "Planification des trajets — création et suivi des trajets",
    "Carte en temps réel — positions GPS de tous les bus actifs avec polylignes de routes",
    "Paiements chauffeurs — gestion des salaires et primes",
    "Notifications — envoi broadcast et individuel",
    "Signalements — consultation et traitement des incidents",
    "Wallet citoyens — validation des demandes de recharge",
    "Alertes système — monitoring des incidents critiques",
]
for page in admin_pages:
    doc.add_paragraph(page, style='List Bullet')

add_placeholder_box(doc, "[ Figure 4 : Interface d'administration — Tableau de bord ]")
add_placeholder_box(doc, "[ Figure 5 : Interface d'administration — Carte GPS en temps réel ]")

doc.add_heading('4.4  Application mobile Chauffeur', level=2)
doc.add_paragraph(
    "L'application mobile chauffeur est développée avec React Native et le framework "
    "Expo SDK 54. Elle est compatible iOS et Android. La navigation est gérée par "
    "React Navigation avec une Stack Navigator pour les écrans d'authentification "
    "et un Bottom Tab Navigator pour les écrans principaux."
)

driver_screens = [
    "SplashScreen — écran de chargement avec vérification du token",
    "LoginScreen — authentification par matricule et mot de passe",
    "DashboardScreen — statut en ligne/hors ligne, bus assigné, trajets planifiés, GPS automatique",
    "TripScreen — gestion du trajet actif (démarrer, pause, terminer), progression des arrêts, envoi GPS continu",
    "ScanTicketScreen — scanner QR des tickets citoyens via la caméra",
    "ReportScreen — formulaire de signalement d'incidents",
    "NotificationsScreen — notifications reçues",
    "PaymentsScreen — historique des paiements",
    "TripsHistoryScreen — historique des trajets effectués",
    "ProfileScreen — informations personnelles, déconnexion",
]
for s in driver_screens:
    doc.add_paragraph(s, style='List Bullet')

add_placeholder_box(doc, "[ Figure 7 : Application chauffeur — Tableau de bord ]")
add_placeholder_box(doc, "[ Figure 8 : Application chauffeur — Trajet actif ]")

doc.add_heading('4.5  Application mobile Citoyen', level=2)
doc.add_paragraph(
    "L'application mobile citoyenne suit le même paradigme technique que l'application "
    "chauffeur, mais avec une couleur primaire verte (#059669) pour différencier "
    "visuellement les deux applications. Elle propose :"
)

citizen_screens = [
    "SplashScreen — animation de démarrage",
    "ChoiceScreen — choix entre Se connecter et Créer un compte",
    "RegisterScreen — inscription (nom, prénom, NNI, téléphone, email, mot de passe)",
    "LoginScreen — authentification par téléphone et mot de passe",
    "DashboardScreen — carte interactive avec bus en temps réel et bus proches",
    "LinesScreen — liste de toutes les lignes avec tarifs et fréquences",
    "LineDetailScreen — détail d'une ligne avec arrêts et bouton de paiement",
    "WalletScreen — solde, historique des transactions, recharge via Bankily",
    "PaymentScreen — confirmation de paiement, choix de méthode",
    "HistoryScreen — historique des trajets avec tickets QR",
    "NotificationsScreen — notifications broadcast",
    "ProfileScreen — informations personnelles, déconnexion",
    "SupportScreen — FAQ et contacts",
    "SettingsScreen — préférences de l'application",
]
for s in citizen_screens:
    doc.add_paragraph(s, style='List Bullet')

add_placeholder_box(doc, "[ Figure 10 : Application citoyenne — Tableau de bord ]")
add_placeholder_box(doc, "[ Figure 11 : Application citoyenne — Écran wallet ]")

doc.add_heading('4.6  Système GPS et carte en temps réel', level=2)
doc.add_paragraph(
    "Le suivi GPS est un composant central de MauriBus. Il fonctionne de la manière "
    "suivante :"
)
gps_flow = [
    "L'application chauffeur utilise expo-location pour obtenir la position GPS du téléphone.",
    "En mode tableau de bord, une position est envoyée toutes les 30 secondes via POST /api/v1/driver/gps/.",
    "En mode trajet actif, la position est mise à jour toutes les 4 secondes (watchPositionAsync).",
    "Le backend stocke chaque position dans le modèle GPSPosition, liée au bus et au trajet actif.",
    "L'interface admin interroge GET /api/v1/admin/live-map/ toutes les 15 secondes.",
    "La carte affiche les marqueurs des bus actifs et les polylignes des routes reliant les arrêts.",
]
for f in gps_flow:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('4.7  Wallet électronique et paiements', level=2)
doc.add_paragraph(
    "Le système de wallet électronique permet aux citoyens de prépayer leurs trajets "
    "en rechargeant leur compte via Bankily (la principale solution de mobile money "
    "en Mauritanie). Le flux de recharge est :"
)
wallet_flow = [
    "Le citoyen envoie de l'argent au numéro Bankily de MauriBus.",
    "Il soumet une demande de recharge dans l'application avec le montant, la référence Bankily et une capture d'écran.",
    "La demande crée une WalletTransaction avec statut EN_ATTENTE.",
    "L'administrateur valide la demande dans l'interface d'administration.",
    "Lors de la validation, le backend crédite automatiquement le WalletCitizen et met à jour la transaction.",
    "Le citoyen peut ensuite utiliser son solde pour payer ses trajets.",
]
for f in wallet_flow:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph(
    "Pour payer un trajet, le citoyen sélectionne une ligne, confirme le paiement, "
    "et le montant correspondant au tarif de la ligne est débité de son wallet. "
    "Un CitizenTrip est créé avec un ticket_code UUID unique qui génère un QR code "
    "dans l'application."
)

doc.add_heading('4.8  Système de tickets QR', level=2)
doc.add_paragraph(
    "Le système de tickets QR permet une validation rapide et sécurisée des billets :"
)
qr_flow = [
    "À l'achat, un CitizenTrip est créé avec un ticket_code UUID unique.",
    "L'application citoyenne affiche ce code sous forme de QR code (react-native-qrcode-svg).",
    "Le chauffeur scanne le QR code via son application (expo-camera + BarCodeScanner).",
    "Le backend vérifie le ticket_code, contrôle que le ticket est PAYE et non UTILISE.",
    "Si valide, le statut est mis à UTILISE avec la date d'utilisation et le trajet associé.",
    "Le chauffeur reçoit une confirmation visuelle avec les informations du passager.",
]
for f in qr_flow:
    doc.add_paragraph(f, style='List Bullet')

add_placeholder_box(doc, "[ Figure 9 : Application chauffeur — Scanner QR ticket ]")
add_placeholder_box(doc, "[ Figure 12 : Application citoyenne — Ticket QR ]")

doc.add_page_break()

# ─── Chapitre 5 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 5 : Tests et Validation', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('5.1  Tests fonctionnels', level=2)
doc.add_paragraph(
    "Les tests fonctionnels ont été réalisés manuellement en suivant les scénarios "
    "d'utilisation définis dans le cahier des charges. Chaque fonctionnalité principale "
    "a été testée en conditions réelles."
)

# Test results table
test_table = doc.add_table(rows=12, cols=4)
test_table.style = 'Table Grid'
test_headers = ['Fonctionnalité', 'Scénario testé', 'Résultat', 'Statut']
for i, h_text in enumerate(test_headers):
    cell = test_table.cell(0, i)
    cell.text = h_text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, 'D1FAE5')

test_data = [
    ['Authentification Admin', 'Login avec matricule/mot de passe', 'JWT retourné', '✓ OK'],
    ['Authentification Chauffeur', 'Login avec bus assigné', 'Bus info retourné', '✓ OK'],
    ['Inscription Citoyen', 'Création compte + wallet', 'Wallet créé automatiquement', '✓ OK'],
    ['Suivi GPS', 'Chauffeur envoie position', 'Position visible sur carte admin', '✓ OK'],
    ['Paiement trajet', 'Citoyen paie ligne à 30 MRU', 'Wallet débité, ticket créé', '✓ OK'],
    ['QR ticket', 'Chauffeur scanne ticket citoyen', 'Ticket validé, statut UTILISE', '✓ OK'],
    ['Recharge wallet', 'Admin valide recharge 200 MRU', 'Wallet crédité de 200 MRU', '✓ OK'],
    ['Notification broadcast', 'Admin envoie notification', 'Visible sur app chauffeur', '✓ OK'],
    ['Signalement', 'Chauffeur crée incident', 'Visible dans admin', '✓ OK'],
    ['Note chauffeur', 'Citoyen note 4/5', 'Moyenne mise à jour', '✓ OK'],
    ['Carte temps réel', 'Bus actif affiché', 'Marqueur + polyligne route', '✓ OK'],
]
for i, row_data in enumerate(test_data):
    row = test_table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

doc.add_heading('5.2  Tests des API', level=2)
doc.add_paragraph(
    "Les endpoints API ont été testés avec Postman en vérifiant les cas normaux, "
    "les cas limites et les cas d'erreur. Les points clés vérifiés :"
)
api_tests = [
    "Authentification : token invalide → 401, token expiré → 401, accès sans token → 403.",
    "Paiement : solde insuffisant → 400 avec message d'erreur explicite.",
    "Validation ticket : ticket déjà utilisé → 400 avec date d'utilisation.",
    "GPS : latitude/longitude manquants → 400.",
    "Recharge : validation double → idempotence vérifiée.",
]
for t in api_tests:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('5.3  Tests des applications mobiles', level=2)
doc.add_paragraph(
    "Les applications mobiles ont été testées via Expo Go sur des appareils physiques "
    "Android. Les tests ont couvert :"
)
mobile_tests = [
    "Flux complet d'inscription et de connexion citoyen.",
    "Flux de paiement de trajet et affichage du QR ticket.",
    "Flux de recharge wallet avec envoi de photo.",
    "Démarrage de trajet chauffeur avec GPS actif.",
    "Scan de ticket QR avec la caméra.",
    "Affichage de la carte avec les bus en temps réel.",
    "Rafraîchissement automatique des données (pull-to-refresh).",
    "Persistance de la session après fermeture de l'application.",
]
for t in mobile_tests:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ─── Chapitre 6 ────────────────────────────────────────────────────────────────
h = doc.add_heading('Chapitre 6 : Résultats et Discussion', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('6.1  Présentation de l\'application', level=2)
doc.add_paragraph(
    "MauriBus constitue une solution opérationnelle complète pour la gestion du transport "
    "urbain. L'application démontre la faisabilité technique d'une telle solution pour "
    "le contexte mauritanien. Les captures d'écran ci-dessous illustrent les principales "
    "fonctionnalités implémentées."
)

add_placeholder_box(doc, "[ Insérer captures d'écran des interfaces ici ]")

doc.add_heading('Métriques techniques', level=3)
metrics = [
    "Plus de 40 endpoints API REST documentés et testés.",
    "16 modèles de données couvrant tous les domaines fonctionnels.",
    "3 applications distinctes partageant un backend commun.",
    "Temps de réponse moyen des API : < 200ms (en développement local).",
    "Suivi GPS avec précision < 10 mètres (selon le matériel).",
    "Authentification JWT avec expiration de 7 jours.",
]
for m in metrics:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('6.2  Difficultés rencontrées', level=2)
doc.add_paragraph(
    "La réalisation de ce projet a présenté plusieurs défis techniques :"
)
difficulties = [
    "Authentification multi-utilisateurs : Django est conçu pour un seul type d'utilisateur. "
    "La gestion de trois types distincts (Admin, Driver, Citizen) a nécessité le développement "
    "de classes d'authentification JWT personnalisées.",
    "Suivi GPS en temps réel : la synchronisation entre les positions envoyées par les "
    "chauffeurs et leur affichage sur la carte admin a nécessité une gestion fine des "
    "requêtes et du polling.",
    "Paiements et précision décimale : l'utilisation de floats Python pour les calculs "
    "financiers peut introduire des erreurs d'arrondi. Nous avons dû passer au type "
    "Decimal pour toutes les opérations sur les wallets.",
    "Permissions React Native : l'accès à la caméra, à la localisation GPS et à la "
    "galerie photo nécessite une gestion explicite des permissions sur Android et iOS.",
    "Communication temps réel : en l'absence de WebSocket (pour simplifier l'architecture), "
    "nous avons utilisé du polling avec des intervalles adaptés selon le contexte (30s "
    "pour le GPS passif, 4s pour le GPS actif en trajet, 15s pour la carte admin).",
]
for i, d in enumerate(difficulties):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d)

doc.add_heading('6.3  Perspectives d\'amélioration', level=2)
doc.add_paragraph(
    "Plusieurs axes d'amélioration sont envisagés pour les versions futures de MauriBus :"
)
improvements = [
    "WebSocket / Server-Sent Events pour un suivi GPS vraiment temps réel sans polling.",
    "Intégration directe avec l'API Bankily pour automatiser la validation des recharges.",
    "Système de notifications push avec Expo Push Notifications via un service dédié.",
    "Algorithme de calcul d'itinéraire multi-lignes pour guider les citoyens.",
    "Tableau de bord analytique avancé avec graphiques d'évolution (revenus, passagers, ponctualité).",
    "Mode hors-ligne pour les applications mobiles (tickets QR disponibles sans connexion).",
    "Migration vers PostgreSQL et déploiement sur serveur cloud (AWS / Azure).",
    "Internationalisation complète (arabe, français) pour une adoption plus large.",
    "Application web progressive (PWA) pour les citoyens ne souhaitant pas installer l'app.",
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_page_break()

# ─── Conclusion ────────────────────────────────────────────────────────────────
h = doc.add_heading('Conclusion Générale', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_paragraph(
    "Ce projet de fin d'études nous a permis de concevoir et de développer MauriBus, "
    "une solution numérique complète pour la gestion du transport urbain en Mauritanie. "
    "En mobilisant des technologies modernes — Django REST Framework, React.js et "
    "React Native — nous avons réussi à créer un écosystème fonctionnel réunissant "
    "trois applications interconnectées au service d'un même objectif : améliorer "
    "la mobilité urbaine à Nouakchott."
)
doc.add_paragraph(
    "Sur le plan technique, ce projet nous a confrontés à des problématiques concrètes "
    "de développement full-stack : conception d'API REST sécurisées, gestion de "
    "l'authentification multi-rôles, traitement de données géospatiales en temps réel, "
    "implémentation d'un système de paiement électronique et développement d'applications "
    "mobiles cross-platform. Ces défis ont été autant d'opportunités d'apprentissage "
    "qui ont renforcé nos compétences en génie logiciel."
)
doc.add_paragraph(
    "Sur le plan fonctionnel, MauriBus répond aux besoins identifiés lors de l'analyse : "
    "les administrateurs disposent d'une vision globale et en temps réel de leur réseau, "
    "les chauffeurs ont accès à des outils modernes de gestion de leurs activités, "
    "et les citoyens bénéficient d'une application intuitive pour accéder aux services "
    "de transport et effectuer leurs paiements."
)
doc.add_paragraph(
    "Les tests effectués ont confirmé la robustesse et la cohérence du système. "
    "Toutes les fonctionnalités principales ont été validées dans des conditions "
    "représentatives de l'utilisation réelle. Le projet constitue une base solide "
    "sur laquelle des améliorations futures pourront être construites."
)
doc.add_paragraph(
    "En définitive, MauriBus démontre qu'il est possible de développer des solutions "
    "numériques de qualité adaptées au contexte africain, avec des moyens limités mais "
    "un engagement fort et les bonnes technologies. Nous espérons que ce travail "
    "contribuera, à sa mesure, à la modernisation du transport urbain mauritanien "
    "et servira de référence pour des projets similaires dans la région."
)

doc.add_page_break()

# ─── Bibliographie ─────────────────────────────────────────────────────────────
h = doc.add_heading('Bibliographie / Webographie', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('Ouvrages et articles', level=2)
refs_books = [
    "Django Software Foundation. (2024). Django Documentation v4.2. djangoproject.com",
    "Feldroy, D. & Greenfeld, A. (2023). Two Scoops of Django 3.x. Feldroy.",
    "Banks, A. & Porcello, E. (2020). Learning React: Modern Patterns for Developing React Apps. O'Reilly.",
    "Boduch, A. & Senra, R. (2021). React and React Native (4th ed.). Packt Publishing.",
]
for r in refs_books:
    p = doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Ressources en ligne', level=2)
refs_web = [
    "Django REST Framework — https://www.django-rest-framework.org/",
    "React Native Documentation — https://reactnative.dev/docs/getting-started",
    "Expo Documentation — https://docs.expo.dev/",
    "React Navigation — https://reactnavigation.org/docs/getting-started/",
    "PyJWT Documentation — https://pyjwt.readthedocs.io/",
    "React Leaflet — https://react-leaflet.js.org/",
    "OpenStreetMap — https://www.openstreetmap.org/",
    "Bankily Mauritanie — https://bankily.mr/",
    "Expo Location API — https://docs.expo.dev/versions/latest/sdk/location/",
    "React Native QR Code SVG — https://github.com/awesomejerry/react-native-qrcode-svg",
]
for r in refs_web:
    p = doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ─── Annexes ───────────────────────────────────────────────────────────────────
h = doc.add_heading('Annexes', level=1)
h.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

doc.add_heading('Annexe A — Captures d\'écran de l\'application', level=2)
add_placeholder_box(doc, "[ Interface d'administration — Tableau de bord ]")
add_placeholder_box(doc, "[ Interface d'administration — Gestion des lignes ]")
add_placeholder_box(doc, "[ Interface d'administration — Carte GPS en temps réel ]")
add_placeholder_box(doc, "[ Application chauffeur — Écran principal ]")
add_placeholder_box(doc, "[ Application citoyenne — Écran principal ]")
add_placeholder_box(doc, "[ Application citoyenne — Wallet ]")

doc.add_heading('Annexe B — Diagrammes UML', level=2)
add_placeholder_box(doc, "[ Diagramme de classes UML ]")
add_placeholder_box(doc, "[ Diagramme de cas d'utilisation ]")

doc.add_heading("Annexe C — Extrait du code source", level=2)
doc.add_paragraph(
    "Le code source complet est disponible dans le dossier du projet. "
    "Ci-dessous un extrait de l'endpoint de validation de ticket QR :"
)

p = doc.add_paragraph()
run = p.add_run(
    "@api_view(['POST'])\n"
    "@authentication_classes([DriverJWTAuthentication])\n"
    "def driver_validate_ticket(request):\n"
    "    ticket_code = request.data.get('ticket_code', '').strip()\n"
    "    ct = CitizenTrip.objects.get(ticket_code=ticket_code)\n"
    "    if ct.statut == 'UTILISE':\n"
    "        return Response({'valid': False, 'detail': 'Ticket déjà utilisé'}, status=400)\n"
    "    ct.statut = 'UTILISE'\n"
    "    ct.date_utilisation = timezone.now()\n"
    "    ct.save()\n"
    "    return Response({'valid': True, 'detail': 'Ticket validé avec succès'})"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ─── Save ───────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'Rapport_MauriBus.docx')
doc.save(output_path)
print(f"[OK] Rapport genere : {output_path}")
